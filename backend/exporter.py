from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from docx import Document
import io

def export_to_pdf(papers: list) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "ResearchIQ - Paper Summaries")
    y -= 40

    for paper in papers:
        if y < 150:
            c.showPage()
            y = height - 50

        c.setFont("Helvetica-Bold", 12)
        title_lines = simpleSplit(paper.get("title", ""), "Helvetica-Bold", 12, width - 100)
        for line in title_lines:
            c.drawString(50, y, line)
            y -= 18

        c.setFont("Helvetica", 10)
        c.drawString(50, y, f"Source: {paper.get('source', '')} | Published: {paper.get('published', '')}")
        y -= 15
        c.drawString(50, y, f"Authors: {paper.get('authors', '')[:80]}")
        y -= 20

        summary = paper.get("summary") or paper.get("abstract", "")
        summary_lines = simpleSplit(summary[:500], "Helvetica", 10, width - 100)
        for line in summary_lines:
            c.drawString(50, y, line)
            y -= 14
        y -= 20

    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def export_to_word(papers: list) -> bytes:
    doc = Document()
    doc.add_heading("ResearchIQ - Paper Summaries", 0)

    for paper in papers:
        doc.add_heading(paper.get("title", "Untitled"), level=1)
        doc.add_paragraph(f"Source: {paper.get('source', '')} | Published: {paper.get('published', '')}")
        doc.add_paragraph(f"Authors: {paper.get('authors', '')}")
        summary = paper.get("summary") or paper.get("abstract", "")
        doc.add_paragraph(summary[:500])
        doc.add_paragraph("---")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()